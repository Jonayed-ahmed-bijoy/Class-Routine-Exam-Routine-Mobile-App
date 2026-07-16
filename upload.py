import os
import re
import uuid
import hashlib
import traceback
import pandas as pd
import requests

from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document

# ==========================================================
# Flask App
# ==========================================================

app = Flask(__name__)

CORS(
    app,
    resources={
        r"/*": {
            "origins": "*"
        }
    }
)

# ==========================================================
# CONFIGURATION
# ==========================================================

FIREBASE_URL = "https://class-routine-app-f3871-default-rtdb.asia-southeast1.firebasedatabase.app/current_routine.json"

EXAM_FIREBASE_URL = "https://class-routine-app-f3871-default-rtdb.asia-southeast1.firebasedatabase.app/current_exam_schedule.json"

# Base (no path) — used to build per-user paths like /users/{uid}/routine.json
FIREBASE_BASE = "https://class-routine-app-f3871-default-rtdb.asia-southeast1.firebasedatabase.app"

# ----------------------------
# Admin Password
#
# Development:
#   admin123
#
# Production:
#   Environment Variable ব্যবহার করবে
# ----------------------------

ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "admin123")

# Optional
FIREBASE_SECRET = os.getenv("FIREBASE_SECRET")

print("=" * 60)
print(" EDU Routine Backend Started")
print("=" * 60)
print(f"Firebase URL : {FIREBASE_URL}")
print(f"Password     : {ADMIN_PASSWORD}")
print(
    "Firebase Auth: Enabled"
    if FIREBASE_SECRET
    else "Firebase Auth: Disabled"
)
print("=" * 60)

# ==========================================================
# Helper Functions
# ==========================================================

VALID_DAYS = [
    "Saturday",
    "Sunday",
    "Monday",
    "Tuesday",
    "Wednesday",
    "Thursday",
    "Friday"
]


def is_time_slot(text):
    """
    Detects whether a cell contains
    a valid time slot.
    """

    if not text:
        return False

    clean_text = (
        text.replace(" ", "")
        .replace("\n", "")
        .replace("\r", "")
    )

    pattern = r"\d{1,2}[\.:]\d{2}[-–—]\d{1,2}[\.:]\d{2}"

    return bool(re.search(pattern, clean_text))


DATE_PATTERN = re.compile(r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4}")


EXAM_TIME_PATTERN = re.compile(
    r"\d{1,2}(?:[\.:]\d{2})?\s*(?:am|pm)?\s*[-–—]\s*"
    r"\d{1,2}(?:[\.:]\d{2})?\s*(?:am|pm)?",
    re.IGNORECASE
)


def is_exam_time_slot(text):
    """
    More lenient than is_time_slot(): exam sheets often
    use whole-hour times like '10 am - 11.30 am' with no
    minutes on the start time.
    """

    if not text:
        return False

    return bool(EXAM_TIME_PATTERN.search(text))


def normalize_exam_time(text):
    match = EXAM_TIME_PATTERN.search(text)
    return clean_text(match.group(0)) if match else clean_text(text)


def is_date(text):
    """
    Detects whether a cell contains a
    DD/MM/YYYY (or similar) date.
    """

    if not text:
        return False

    return bool(DATE_PATTERN.search(text))


def extract_day(text):
    """
    Pulls a weekday name out of a cell that may
    contain both a date and a day, e.g.
    '11/07/2026 Saturday'.
    """

    if not text:
        return None

    for day in VALID_DAYS:
        if day.lower() in text.lower():
            return day

    return None


def clean_text(text):
    """
    Removes unnecessary spaces/newlines.
    """

    if text is None:
        return ""

    return (
        text.replace("\n", " ")
        .replace("\r", " ")
        .strip()
    )


def build_firebase_url(base_url=FIREBASE_URL):
    """
    Returns Firebase URL.
    If FIREBASE_SECRET exists,
    append ?auth=
    """

    if FIREBASE_SECRET:
        return f"{base_url}?auth={FIREBASE_SECRET}"

    return base_url


def log(message):
    """
    Simple logger
    """

    print(f"[INFO] {message}")


def log_error(message):
    """
    Error logger
    """

    print(f"[ERROR] {message}")

# ==========================================================
# Faculty Mapping
# ==========================================================

def get_faculty_mapping(doc):
    """
    Reads the faculty table from the DOCX and returns

    {
        "ABC": "Dr. ABC XYZ",
        "MRH": "Md Rahim",
        ...
    }
    """

    mapping = {}

    for table in doc.tables:

        rows = table.rows

        if not rows:
            continue

        for i in range(min(3, len(rows))):

            header = [
                clean_text(cell.text).lower()
                for cell in rows[i].cells
            ]

            if (
                "name" in header
                and (
                    "short form" in header
                    or "acronym" in header
                )
            ):

                try:

                    name_idx = header.index("name")

                    acronym_idx = None

                    for idx, h in enumerate(header):

                        if (
                            "short" in h
                            or "acronym" in h
                        ):
                            acronym_idx = idx
                            break

                    if acronym_idx is None:
                        continue

                    for row in rows[i + 1:]:

                        cells = row.cells

                        if len(cells) <= max(name_idx, acronym_idx):
                            continue

                        faculty_name = clean_text(
                            cells[name_idx].text
                        )

                        faculty_short = clean_text(
                            cells[acronym_idx].text
                        )

                        if faculty_name and faculty_short:

                            mapping[
                                faculty_short
                            ] = faculty_name

                except Exception:

                    continue

                break

    log(f"Faculty Loaded : {len(mapping)}")

    return mapping


# ==========================================================
# Document Type Detection
# ==========================================================

def detect_document_type(doc):
    """
    Looks at the header row of every table.

    If a header cell says something like "Date"
    (exam routines have one, class routines never do),
    treat this file as an Exam Schedule.

    Otherwise treat it as a normal Class Routine.
    """

    for table in doc.tables:

        rows = table.rows

        if not rows:
            continue

        for i in range(min(3, len(rows))):

            header = [
                clean_text(cell.text).lower()
                for cell in rows[i].cells
            ]

            if any("date" == h or "date" in h for h in header):
                return "exam_schedule"

    return "class_routine"


# ==========================================================
# Exam Schedule Parser
# ==========================================================

def parse_exam_schedule(doc):
    """
    Parses an Exam Routine DOCX.

    Only Date, Day, Time and Course Code are kept —
    exam sheets don't reliably carry per-course Room/
    Faculty the way class routines do, and the app
    doesn't need them for this table.
    """

    schedule = []

    course_pattern = re.compile(
        r"([A-Z]{2,4}\s*\d{3}(?:\.[0-9A-Za-z]+)?)",
        re.IGNORECASE
    )

    current_date = None
    current_day = None
    current_time = None

    for table in doc.tables:

        rows = table.rows

        if not rows:
            continue

        # ------------------------------------
        # Find header row (contains "date")
        # ------------------------------------

        header_index = -1
        date_idx = time_idx = subject_idx = -1

        for i in range(min(3, len(rows))):

            header = [
                clean_text(cell.text).lower()
                for cell in rows[i].cells
            ]

            if any("date" in h for h in header):

                header_index = i

                for idx, h in enumerate(header):

                    if "date" in h:
                        date_idx = idx
                    elif "time" in h:
                        time_idx = idx
                    elif "subject" in h or "course" in h:
                        subject_idx = idx

                break

        if header_index == -1 or subject_idx == -1:
            continue

        # ------------------------------------
        # Walk data rows (dates/times only
        # appear on the first row of a merged
        # group — carry them forward)
        # ------------------------------------

        for row in rows[header_index + 1:]:

            cells = row.cells

            if len(cells) <= subject_idx:
                continue

            if date_idx != -1 and date_idx < len(cells):

                date_text = clean_text(cells[date_idx].text)

                if is_date(date_text):
                    current_date = DATE_PATTERN.search(date_text).group(0)

                day_found = extract_day(date_text)

                if day_found:
                    current_day = day_found

            if time_idx != -1 and time_idx < len(cells):

                time_text = clean_text(cells[time_idx].text)

                if is_exam_time_slot(time_text):
                    current_time = normalize_exam_time(time_text)

            if current_date is None or current_time is None:
                continue

            subject_text = clean_text(cells[subject_idx].text)

            if not subject_text:
                continue

            courses = course_pattern.findall(subject_text)

            for course in courses:

                schedule.append({
                    "Date": current_date,
                    "Day": current_day or "",
                    "Time": current_time,
                    "Course": clean_text(course),
                })

    log(f"Exam Schedule Parsed : {len(schedule)} Entries")

    dataframe = pd.DataFrame(schedule)

    if not dataframe.empty:
        dataframe = dataframe.drop_duplicates(
            subset=["Date", "Time", "Course"]
        ).reset_index(drop=True)

    log(f"Exam Schedule After Dedup : {len(dataframe)} Entries")

    return dataframe


# ==========================================================
# Routine Parser
# ==========================================================

def parse_routine_complete(doc):

    faculty_map = get_faculty_mapping(doc)

    routine = []

    course_pattern = re.compile(
        r"([A-Z]{2,4}\s*\d{3}(?:\.[0-9A-Za-z]+)?)",
        re.IGNORECASE
    )

    custom_time_pattern = re.compile(
        r"\(?(\d{1,2}[\.:]\d{2}\s*[-–—]\s*\d{1,2}[\.:]\d{2})\)?"
    )

    current_day = None

    for table in doc.tables:

        rows = table.rows

        if not rows:
            continue

        first_text = clean_text(rows[0].cells[0].text)

        if first_text in VALID_DAYS:
            current_day = first_text

        header_index = -1
        time_slots = []

        # ------------------------------------
        # Find Time Slot Row
        # ------------------------------------

        for i in range(min(5, len(rows))):

            cells = rows[i].cells

            slots = []

            for idx, cell in enumerate(cells):

                txt = clean_text(cell.text)

                if is_time_slot(txt):

                    slots.append({
                        "index": idx,
                        "time": txt
                    })

            if slots:

                header_index = i
                time_slots = slots
                break

        if header_index == -1:
            continue

        # ------------------------------------
        # Read Remaining Rows
        # ------------------------------------

        for row in rows[header_index + 1:]:

            cells = row.cells

            if not cells:
                continue

            first_cell = clean_text(cells[0].text)

            if first_cell in VALID_DAYS:

                current_day = first_cell
                continue

            if current_day is None:
                continue

            room = first_cell

            if not (
                room.isdigit()
                or room.upper().startswith("N")
            ):
                continue

            for slot in time_slots:

                idx = slot["index"]

                if idx >= len(cells):
                    continue

                raw_text = clean_text(
                    cells[idx].text
                )

                if not raw_text:
                    continue

                processing = raw_text

                final_time = slot["time"]

                # ------------------------------------
                # Custom Time
                # ------------------------------------

                match = custom_time_pattern.search(
                    processing
                )

                if match:

                    final_time = (
                        match.group(1)
                        .replace(" ", "")
                    )

                    processing = processing.replace(
                        match.group(0),
                        ""
                    )

                # ------------------------------------
                # Course Code
                # ------------------------------------

                course = course_pattern.search(
                    processing
                )

                if not course:
                    continue

                course_code = clean_text(
                    course.group(1)
                )

                remaining = processing.replace(
                    course_code,
                    ""
                )

                remaining = re.sub(
                    r"\([A-Z]{2,5}\)",
                    "",
                    remaining
                )

                faculty = remaining.strip(
                    " ,-–—()"
                )

                if len(faculty) > 15:
                    faculty = ""

                header_text = " ".join([
                    clean_text(c.text).lower()
                    for c in rows[header_index].cells
                ])

                previous = ""

                if header_index > 0:

                    previous = clean_text(
                        rows[
                            header_index - 1
                        ].cells[0].text
                    ).lower()

                class_type = (
                    "Lab"
                    if (
                        "lab" in header_text
                        or "lab" in previous
                    )
                    else "Theory"
                )

                routine.append({

                    "Day": current_day,

                    "Time": final_time,

                    "Room": room,

                    "Course": course_code,

                    "FacultyAcronym": faculty,

                    "FacultyFullName":
                        faculty_map.get(
                            faculty,
                            ""
                        ),

                    "Type": class_type
                })

    log(f"Routine Parsed : {len(routine)} Classes")

    return pd.DataFrame(routine)
# ==========================================================
# ROUTES
# ==========================================================

@app.route("/")
def home():

    return jsonify({
        "success": True,
        "message": "EDU Routine Backend is Running",
        "version": "2.0"
    })


@app.route("/health")
def health():

    return jsonify({
        "status": "healthy"
    })


# ==========================================================
# User Accounts (simple, self-hosted — see note in DEPLOY docs)
# ==========================================================
#
# Each user is stored at /users/{uid}.json = {email, password_hash,
# created_at}. A lookup index at /email_index/{email_hash}.json = uid
# lets us find a uid from an email at login time without scanning all
# users. Every uploaded routine/exam-schedule is then stored under
# /users/{uid}/routine.json or /users/{uid}/exam_schedule.json instead
# of the old shared /current_routine.json — so it never shows up for
# anyone else.

def firebase_url_for(path):
    suffix = f"?auth={FIREBASE_SECRET}" if FIREBASE_SECRET else ""
    return f"{FIREBASE_BASE}/{path}.json{suffix}"


def email_hash(email):
    return hashlib.sha256(email.strip().lower().encode()).hexdigest()


def find_uid_by_email(email):
    resp = requests.get(
        firebase_url_for(f"email_index/{email_hash(email)}"),
        timeout=15,
    )
    if resp.status_code != 200:
        return None
    return resp.json()


def get_user(uid):
    resp = requests.get(firebase_url_for(f"users/{uid}"), timeout=15)
    if resp.status_code != 200:
        return None
    return resp.json()


EMAIL_PATTERN = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")


@app.route("/api/signup", methods=["POST"])
def signup():
    try:
        body = request.get_json(force=True, silent=True) or {}

        email = (body.get("email") or "").strip().lower()
        password = body.get("password") or ""

        if not EMAIL_PATTERN.match(email):
            return jsonify({"success": False, "error": "Invalid email address"}), 400

        if len(password) < 6:
            return jsonify({
                "success": False,
                "error": "Password must be at least 6 characters"
            }), 400

        if find_uid_by_email(email) is not None:
            return jsonify({
                "success": False,
                "error": "An account with this email already exists"
            }), 409

        uid = uuid.uuid4().hex

        user_record = {
            "email": email,
            "password_hash": generate_password_hash(password),
            "created_at": str(pd.Timestamp.now()),
        }

        requests.put(firebase_url_for(f"users/{uid}"), json=user_record, timeout=15)
        requests.put(firebase_url_for(f"email_index/{email_hash(email)}"), json=uid, timeout=15)

        log(f"New signup: {email} ({uid})")

        return jsonify({"success": True, "uid": uid, "email": email})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/login", methods=["POST"])
def login():
    try:
        body = request.get_json(force=True, silent=True) or {}

        email = (body.get("email") or "").strip().lower()
        password = body.get("password") or ""

        uid = find_uid_by_email(email)

        if uid is None:
            return jsonify({"success": False, "error": "No account with this email"}), 404

        user = get_user(uid)

        if user is None or not check_password_hash(user.get("password_hash", ""), password):
            return jsonify({"success": False, "error": "Incorrect password"}), 401

        log(f"Login: {email} ({uid})")

        return jsonify({"success": True, "uid": uid, "email": email})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/upload", methods=["POST"])
def upload_file():

    try:

        log("=" * 50)
        log("New Upload Request Received")

        # ----------------------------------------
        # User Check (replaces the old shared admin
        # password — each upload now belongs to the
        # signed-in user only)
        # ----------------------------------------

        uid = request.form.get("uid", "").strip()

        if not uid or get_user(uid) is None:

            log_error("Invalid Or Missing User Session")

            return jsonify({
                "success": False,
                "error": "You must be signed in to upload"
            }), 401

        log(f"Uploading As User : {uid}")

        # ----------------------------------------
        # File Check
        # ----------------------------------------

        if "file" not in request.files:

            log_error("No File Uploaded")

            return jsonify({
                "success": False,
                "error": "No file uploaded"
            }), 400

        file = request.files["file"]

        if file.filename == "":

            log_error("Empty Filename")

            return jsonify({
                "success": False,
                "error": "No file selected"
            }), 400

        if not file.filename.lower().endswith(".docx"):

            log_error("Invalid File Type")

            return jsonify({
                "success": False,
                "error": "Only DOCX files are allowed"
            }), 400

        log(f"File : {file.filename}")

        

        # ----------------------------------------
        # Parse DOCX
        # ----------------------------------------

        log("Reading DOCX...")

        document = Document(file.stream)

        doc_type = detect_document_type(document)

        log(f"Detected Document Type : {doc_type}")

        if doc_type == "exam_schedule":
            dataframe = parse_exam_schedule(document)
            target_firebase_url = firebase_url_for(f"users/{uid}/exam_schedule")
        else:
            dataframe = parse_routine_complete(document)
            target_firebase_url = firebase_url_for(f"users/{uid}/routine")

        routine_json = dataframe.to_dict(
            orient="records"
        )

        log(
            f"Total Entries : {len(routine_json)}"
        )

        # ----------------------------------------
        # Firebase Upload
        # ----------------------------------------

        payload = {

            "updatedAt": str(pd.Timestamp.now()),

            #"isRamadan": is_ramadan,

            "data": routine_json

        }

        log("Uploading To Firebase...")

        firebase_response = requests.put(

            target_firebase_url,

            json=payload,

            timeout=30

        )

        # ----------------------------------------
        # Firebase Failed
        # ----------------------------------------

        if firebase_response.status_code != 200:

            log_error(
                firebase_response.text
            )

            return jsonify({

                "success": False,

                "error":
                    "Firebase Upload Failed",

                "firebase_status":
                    firebase_response.status_code

            }), 500

        # ----------------------------------------
        # Success
        # ----------------------------------------

        log("Firebase Upload Successful")
        log("=" * 50)

        return jsonify({

            "success": True,

            "message":
                "Exam Schedule Uploaded Successfully"
                if doc_type == "exam_schedule"
                else "Routine Uploaded Successfully",

            "type":
                doc_type,

            "count":
                len(routine_json),

            "firebase_status":
                firebase_response.status_code,

           # "isRamadan":
               # is_ramadan

        })

    except Exception as e:

        traceback.print_exc()

        log_error(str(e))

        return jsonify({

            "success": False,

            "error": str(e)

        }), 500   
# ==========================================================
# MAIN
# ==========================================================

if __name__ == "__main__":

    print("\n")
    print("=" * 60)
    print(" EDU Routine Generator Backend")
    print("=" * 60)
    print(" Server Started Successfully")
    print(" URL : http://127.0.0.1:8000")
    print(" Health : http://127.0.0.1:8000/health")
    print(" Upload : http://127.0.0.1:8000/api/upload")
    print("=" * 60)

    app.run(
        host="0.0.0.0",
        port=8000,
        debug=True
    )     